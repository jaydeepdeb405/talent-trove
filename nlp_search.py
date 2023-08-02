from sentence_transformers import SentenceTransformer, util
from docx import Document
from nltk.corpus import stopwords
from nltk.stem import WordNetLemmatizer
from nltk.tokenize import word_tokenize
from data_service import refresh_candidates, resume_root_path
import os
import string
import nltk
import tensorflow as tf

# Initialize Sentence Transformer model and NLTK objects
model_name = 'all-MiniLM-L6-v2'
model = SentenceTransformer(model_name)
lemmatizer = WordNetLemmatizer()
nltk.download('stopwords')
stopwords_set = set(stopwords.words('english'))

nltk.download('punkt')
nltk.download('wordnet')


# Preprocess text
def preprocess_text(text):
    # Convert to lowercase
    text = text.lower()

    # Remove punctuation
    text = text.translate(str.maketrans('', '', string.punctuation))

    # Tokenization
    tokens = word_tokenize(text)

    # Remove stopwords
    tokens = [token for token in tokens if token not in stopwords_set]

    # Lemmatization
    tokens = [lemmatizer.lemmatize(token) for token in tokens]

    # Return preprocessed text as a string
    return ' '.join(tokens)


def candidate_search(search_text, max_results=5):
    preprocessed_search_text = preprocess_text(search_text)

    # Calculate cosine similarity
    cosine_similarity_scores = {}

    for folder_name, subfolders, filenames in os.walk(resume_root_path):
         for subfolder in subfolders:
            for filename in os.listdir(f"{resume_root_path}/{subfolder}"):
                file_path = f"{resume_root_path}/{subfolder}/{filename}"
                if filename.endswith('.docx'):
                    with open(file_path, 'rb') as f:
                        doc = Document(f)
                        resume_text = ' '.join([p.text for p in doc.paragraphs])
                elif filename.endswith('.txt'):
                    with open(file_path, encoding="utf8") as f:
                        resume_text = f.read()
                preprocessed_resume_text = preprocess_text(resume_text)
                search_text_embedding = model.encode([preprocessed_search_text])
                resume_embedding = model.encode([preprocessed_resume_text])
                cosine_similarity = util.cos_sim(search_text_embedding, resume_embedding)[0][0]
                cosine_similarity_scores[subfolder] = cosine_similarity

    # Sort and retrieve top x matches
    top_matches = sorted(cosine_similarity_scores.items(), key=lambda x: x[1], reverse=True)[:max_results]

    # Print top x matching resumes with file names and cosine similarity scores
    matches = {}
    for match in top_matches:
        file_path, similarity_score = match
        file_name = os.path.basename(file_path)
        matches[file_name] = tf.as_string(similarity_score).numpy().decode('ascii')

    return matches


def resume_search(job_description, email_list=[], max_results=10):
    matches = {}
    cosine_similarity_scores = {}

    preprocessed_job_description = preprocess_text(job_description)

    root_folder = resume_root_path  # Replace with the actual root folder path
    total_emails = len(email_list)
    max_results = total_emails if total_emails > max_results else max_results
    if total_emails > 0:
        for email in email_list:
            file_path = f"{root_folder}/{email}"
            if not os.path.exists(file_path):
                refresh_candidates(email)

            for folder_name, subfolders, filenames in os.walk(file_path):
                for filename in filenames:
                    file_path = os.path.join(folder_name, filename)
                    cosine_similarity_scores[email] = get_cosine_similarity(preprocessed_job_description, file_path)
    else:
        for folder_name, subfolders, filenames in os.walk(root_folder):
            for filename in filenames:
                file_path = os.path.join(folder_name, filename)
                cosine_similarity_scores[folder_name] = get_cosine_similarity(preprocessed_job_description, file_path)

    # Sort and retrieve top 5 matches
    top_matches = sorted(cosine_similarity_scores.items(), key=lambda x: x[1], reverse=True)[:max_results]

    # Print top 5 matching resumes with file names and cosine similarity scores
    for match in top_matches:
        file_path, similarity_score = match
        file_name = os.path.basename(file_path)
        matches[file_name] = tf.as_string(similarity_score).numpy().decode('ascii')

    return matches
    # return matches


def get_cosine_similarity(preprocessed_job_description, file_path):
    filename = os.path.basename(file_path)
    resume_text = ''
    if filename.endswith('.docx'):
        with open(file_path, 'rb') as f:
            doc = Document(f)
            resume_text = ' '.join([p.text for p in doc.paragraphs])
    elif filename.endswith('.txt'):
        with open(file_path, encoding="utf8") as f:
            resume_text = f.read()
    preprocessed_resume_text = preprocess_text(resume_text)
    job_description_embedding = model.encode([preprocessed_job_description])
    resume_embedding = model.encode([preprocessed_resume_text])
    cosine_similarity = util.cos_sim(job_description_embedding, resume_embedding)[0][0]
    return cosine_similarity
